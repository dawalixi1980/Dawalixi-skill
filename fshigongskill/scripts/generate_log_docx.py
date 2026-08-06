#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 Word 施工日志（A4 横向，现代简约排版，套用 word-formatter simple_modern 模板规范）。

排版要点（源自 word-formatter simple_modern 模板）：
- 页面：A4 横向，页边距 20mm，页脚居中页码
- 标题：深蓝色加粗居中 + 底部强调线
- 信息表：标签列浅蓝底加粗，值列常规
- 施工记录表：表头深蓝底白字，奇数行浅灰斑马纹，浅灰细边框，单元格垂直居中、内边距适中
- 附记：分节小标题蓝色加粗
- 照片附图：两列表格网格排版，照片居中带浅灰边框，图注 9pt 灰色居中

用法：
    python scripts/generate_log_docx.py --project "项目名称" --date 2026-07-14 --data 日志数据.json [--base-dir 目录]

日志数据.json 结构：
{
  "工程名称": "XX工程",
  "施工单位": "XX公司",
  "天气": "晴 25-32℃",
  "记录人": "张三",
  "施工日志": [
    {
      "时间": "09:00-11:00",
      "标段": "A标段",
      "桩号/部位": "1号桩",
      "工序": "浇筑",
      "施工内容": "...",
      "班组/机械/材料": "...",
      "验收情况": "...",
      "照片": ["2026-07-14_A标段_1号桩_浇筑.jpg"]
    }
  ],
  "质量安全及问题处理": "...",
  "设计变更与技术交底": "...",
  "各方协调沟通": "...",
  "次日施工计划": "..."
}

照片文件按文件名在项目"照片归档/"下递归查找。
"""

import argparse
import json
import os
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
HEADING_BLUE = RGBColor(0x19, 0x76, 0xD2)   # #1976D2 现代蓝
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)        # #333333 深灰正文
HEADER_FILL = "1976D2"                        # 表头深蓝底
HEADER_TEXT = "FFFFFF"                        # 表头白字
ZEBRA_FILL = "F2F2F2"                         # 斑马纹浅灰
LABEL_FILL = "E8F0FE"                         # 信息表标签浅蓝底
GRAY_BORDER = "BFBFBF"                        # 浅灰边框
PHOTO_GRID_COLS = 2                           # 照片网格列数

LOG_HEADERS = ["序号", "时间", "标段", "桩号/部位", "工序",
               "施工内容", "班组/机械/材料", "验收情况", "对应照片"]
LOG_WIDTHS = [1.0, 1.6, 1.8, 2.0, 2.0, 6.2, 3.4, 2.8, 3.0]


def set_run(run, size=10.5, bold=False, color=DARK_TEXT):
    run.font.name = EN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CN_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def set_table_borders(table, color=GRAY_BORDER, sz="4"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def set_cell(cell, text, bold=False, size=10.5, center=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        if line:
            set_run(p.add_run(line), size=size, bold=bold)
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)


def add_bottom_border(paragraph, color=GRAY_BORDER, sz="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text), size=20, bold=True, color=HEADING_BLUE)
    add_bottom_border(p, color=HEADER_FILL, sz="16")


def add_section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text), size=13, bold=True, color=HEADING_BLUE)


def add_info_table(doc, pairs, label_w=3.0, value_w=6.5):
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for (l1, v1), (l2, v2) in pairs:
        row = table.add_row()
        c = row.cells
        set_cell(c[0], l1, bold=True, center=True, fill=LABEL_FILL)
        set_cell(c[1], v1)
        set_cell(c[2], l2, bold=True, center=True, fill=LABEL_FILL)
        set_cell(c[3], v2)
        c[0].width = Cm(label_w)
        c[1].width = Cm(value_w)
        c[2].width = Cm(label_w)
        c[3].width = Cm(value_w)
    return table


def add_log_table(doc, entries):
    table = doc.add_table(rows=1, cols=len(LOG_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    for i, h in enumerate(LOG_HEADERS):
        set_cell(hdr.cells[i], h, bold=True, center=True, fill=HEADER_FILL, size=10.5)
        hdr.cells[i].width = Cm(LOG_WIDTHS[i])
    for idx, e in enumerate(entries, 1):
        row = table.add_row()
        vals = [
            str(idx),
            e.get("时间", ""),
            e.get("标段", ""),
            e.get("桩号/部位", ""),
            e.get("工序", ""),
            e.get("施工内容", ""),
            e.get("班组/机械/材料", ""),
            e.get("验收情况", ""),
            "\n".join(e.get("照片", [])),
        ]
        fill = ZEBRA_FILL if idx % 2 == 0 else None
        for i, v in enumerate(vals):
            set_cell(row.cells[i], v, fill=fill, size=10)
            row.cells[i].width = Cm(LOG_WIDTHS[i])
    return table


def add_note(doc, label, content):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    set_run(p.add_run(label + "："), size=11, bold=True, color=HEADING_BLUE)
    if content:
        set_run(p.add_run(content), size=11)


def add_sign_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    set_run(p.add_run("记录人签字：____________________　项目技术负责人签字：____________________"),
            size=11)


def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("第 ")
        set_run(run, size=9, color=RGBColor(0x80, 0x80, 0x80))
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), " PAGE ")
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rpr.append(sz)
        r.append(rpr)
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        p._p.append(fld)
        run2 = p.add_run(" 页")
        set_run(run2, size=9, color=RGBColor(0x80, 0x80, 0x80))


def find_photo(root, fname):
    for dirpath, _, files in os.walk(root):
        if fname in files:
            return os.path.join(dirpath, fname)
    return None


def add_photo_appendix(doc, entries, root):
    photos = []
    for e in entries:
        for fname in e.get("照片", []):
            p = find_photo(root, fname)
            if p:
                photos.append((fname, p))
    if not photos:
        doc.add_paragraph("（照片归档在 照片归档/ 目录，详见项目-Wiki）")
        return 0

    n_cols = PHOTO_GRID_COLS
    n_rows = (len(photos) + n_cols - 1) // n_cols
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF", sz="0")

    total_w = 25.0
    col_w = total_w / n_cols
    for i, (fname, p) in enumerate(photos):
        r, c = divmod(i, n_cols)
        cell = table.cell(r, c)
        cell.width = Cm(col_w)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(p, width=Cm(col_w - 2.0))
        except Exception as ex:
            print(f"  跳过无法插入的图片：{fname}（{ex}）")
            para.text = ""
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(cap.add_run(fname), size=9, color=RGBColor(0x80, 0x80, 0x80))
    return len(photos)


def load_json(path):
    for enc in ("utf-8-sig", "gbk", "utf-8"):
        try:
            with open(path, "r", encoding=enc) as f:
                return json.load(f)
        except (UnicodeDecodeError, json.JSONDecodeError):
            continue
    raise ValueError(f"无法解析日志数据 JSON：{path}")


def main():
    parser = argparse.ArgumentParser(description="生成 Word 施工日志（现代简约排版）")
    parser.add_argument("--project", required=True, help="项目名称（也是项目根目录名）")
    parser.add_argument("--date", required=True, help="日志日期 YYYY-MM-DD")
    parser.add_argument("--data", required=True, help="日志数据 JSON 路径")
    parser.add_argument("--base-dir", default=".", help="项目根目录的父目录（默认当前目录）")
    args = parser.parse_args()

    root = os.path.join(args.base_dir, args.project)
    data = load_json(args.data)

    doc = Document()

    # A4 横向，20mm 页边距
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    add_title(doc, "施 工 日 志")

    info_rows = [
        (("工程名称", data.get("工程名称") or args.project), ("施工单位", data.get("施工单位", ""))),
        (("日期", args.date), ("天气", data.get("天气", ""))),
    ]
    if data.get("建设单位"):
        info_rows.append((("建设单位", data["建设单位"]), ("监理单位", data.get("监理单位", ""))))
    info_rows.append((("记录人", data.get("记录人", "")), ("项目技术负责人", "")))
    add_info_table(doc, info_rows)

    add_section(doc, "一、施工记录")
    entries = data.get("施工日志", [])
    if entries:
        add_log_table(doc, entries)
    else:
        doc.add_paragraph("（当日无施工作业，本栏留空）")

    add_section(doc, "二、附记")
    add_note(doc, "质量安全及问题处理", data.get("质量安全及问题处理", ""))
    add_note(doc, "设计变更与技术交底", data.get("设计变更与技术交底", ""))
    add_note(doc, "各方协调沟通", data.get("各方协调沟通", ""))
    add_note(doc, "次日施工计划", data.get("次日施工计划", ""))

    if entries:
        add_section(doc, "三、当日照片附图")
        add_photo_appendix(doc, entries, root)
    else:
        add_section(doc, "三、当日照片附图")
        doc.add_paragraph("（无）")

    add_sign_block(doc)
    add_page_number_footer(doc)

    log_dir = os.path.join(root, "施工日志")
    os.makedirs(log_dir, exist_ok=True)
    out_path = os.path.join(log_dir, f"{args.date}_施工日志.docx")
    doc.save(out_path)
    print(f"施工日志已生成：{out_path}")


if __name__ == "__main__":
    main()
